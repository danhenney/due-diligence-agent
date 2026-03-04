"""DOCX report generator — Word document from DD state."""
from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _safe_str(val: Any, default: str = "N/A") -> str:
    if val is None:
        return default
    if isinstance(val, dict):
        return val.get("summary", val.get("description", str(val)[:300]))
    if isinstance(val, list):
        return "; ".join(str(v) for v in val[:5])
    return str(val)[:500]


def _extract_bullets(data: Any, keys: list[str], max_items: int = 10) -> list[str]:
    """Extract bullet points from a dict, trying multiple key patterns."""
    if not isinstance(data, dict):
        return [_safe_str(data)]
    bullets = []
    for key in keys:
        val = data.get(key)
        if val is None:
            continue
        if isinstance(val, list):
            for item in val[:max_items]:
                if isinstance(item, dict):
                    label = (item.get("name") or item.get("risk") or item.get("trend")
                             or item.get("argument") or item.get("question") or "")
                    desc = (item.get("description") or item.get("impact")
                            or item.get("supporting_evidence") or "")
                    if isinstance(desc, list):
                        desc = "; ".join(str(d) for d in desc[:3])
                    bullets.append(f"{label}: {desc}" if label else str(desc)[:300])
                else:
                    bullets.append(str(item)[:300])
        elif isinstance(val, str):
            bullets.append(val[:500])
        elif isinstance(val, dict):
            summary = (val.get("summary") or val.get("description")
                       or val.get("assessment") or str(val)[:300])
            bullets.append(f"{key}: {summary}")
    if not bullets:
        summary = data.get("summary", "")
        if summary:
            bullets.append(summary)
    return bullets[:max_items]


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_section(doc: Document, title: str, bullets: list[str],
                 subtitle: str = "") -> None:
    _add_heading(doc, title, level=2)
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")


def _add_markdown_report(doc: Document, markdown_text: str) -> None:
    """Parse final_report markdown into Word paragraphs."""
    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Bullets
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        # Horizontal rules
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 50)
        # Regular paragraph (strip bold markdown)
        else:
            cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            doc.add_paragraph(cleaned)


def generate_docx(state: dict[str, Any], job_id: str,
                  output_dir: str | None = None) -> str:
    """Generate a Word document from DD state.

    Returns absolute path to the generated .docx file.
    """
    reports_dir = Path(output_dir) if output_dir else Path("reports")
    reports_dir.mkdir(parents=True, exist_ok=True)
    output_path = str(reports_dir / f"{job_id}.docx")

    company = state.get("company_name") or "Unknown Company"
    recommendation = (state.get("recommendation") or "WATCH").upper()

    doc = Document()

    # ── Cover / Title ────────────────────────────────────────────────────────
    title = doc.add_heading(f"Due Diligence Report: {company}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Recommendation: {recommendation}")
    run.font.size = Pt(20)
    run.bold = True
    if "INVEST" in recommendation:
        run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    elif "PASS" in recommendation:
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    else:
        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── Executive Summary ────────────────────────────────────────────────────
    si = state.get("strategic_insight") or {}
    _add_section(doc, "Executive Summary", [
        (state.get("final_report", "") or "")[:600],
        f"Recommendation: {recommendation}",
        f"Rationale: {_safe_str(si.get('rationale', 'See full report'))}",
    ])

    # ── Market Analysis ──────────────────────────────────────────────────────
    market = state.get("market_analysis") or {}
    _add_section(doc, "Market Analysis",
        _extract_bullets(market, ["summary", "tam", "sam", "som", "cagr", "trends", "market_drivers"]),
        subtitle="TAM / SAM / SOM / Growth Trends")

    # ── Competitive Landscape ────────────────────────────────────────────────
    comp = state.get("competitor_analysis") or {}
    _add_section(doc, "Competitive Landscape",
        _extract_bullets(comp, ["summary", "competitors", "market_share", "competitive_gaps"]),
        subtitle="Competitors & Positioning")

    # ── Financial Analysis ───────────────────────────────────────────────────
    fin = state.get("financial_analysis") or {}
    _add_section(doc, "Financial Analysis & Valuation",
        _extract_bullets(fin, ["summary", "revenue_trend", "profitability", "cash_flow", "valuation", "key_ratios"]),
        subtitle="Revenue / Profitability / Cash Flow / Fair Value")

    # ── Technology & IP ──────────────────────────────────────────────────────
    tech = state.get("tech_analysis") or {}
    _add_section(doc, "Technology & IP",
        _extract_bullets(tech, ["summary", "core_technologies", "ip_patents", "tech_maturity"]),
        subtitle="Core Tech / Patents / Maturity")

    # ── Legal & Regulatory ───────────────────────────────────────────────────
    legal = state.get("legal_regulatory") or {}
    _add_section(doc, "Legal & Regulatory",
        _extract_bullets(legal, ["summary", "investment_structure_risks", "business_regulatory_risks", "litigation"]),
        subtitle="Investment Structure & Business Regulatory Risks")

    # ── Team & Leadership ────────────────────────────────────────────────────
    team = state.get("team_analysis") or {}
    _add_section(doc, "Team & Leadership",
        _extract_bullets(team, ["summary", "leadership_profiles", "capability_assessment", "key_person_risk"]),
        subtitle="Leadership / Capability / Key Person Risk")

    # ── Investment Thesis ────────────────────────────────────────────────────
    ras = state.get("ra_synthesis") or {}
    _add_section(doc, "Investment Thesis",
        _extract_bullets(ras, ["summary", "core_investment_arguments", "attractiveness_scorecard", "key_findings"]),
        subtitle="Core Arguments & Attractiveness Scorecard")

    # ── Risk Assessment ──────────────────────────────────────────────────────
    risk = state.get("risk_assessment") or {}
    _add_section(doc, "Risk Assessment",
        _extract_bullets(risk, ["summary", "top_risks", "risk_matrix", "overall_risk_level"]),
        subtitle="Risk Matrix / Top Risks / Mitigation")

    # ── DD Questions ─────────────────────────────────────────────────────────
    ddq = state.get("dd_questions") or {}
    questions = ddq.get("dd_questionnaire", [])
    q_bullets = []
    for q in questions[:12]:
        if isinstance(q, dict):
            q_bullets.append(f"[{q.get('priority', '?')}] {q.get('question', '?')} — Target: {q.get('target', '?')}")
        else:
            q_bullets.append(str(q)[:300])
    _add_section(doc, "DD Questionnaire",
        q_bullets or ["No unresolved questions identified."],
        subtitle="Outstanding Questions for Follow-Up")

    # ── Full Report ──────────────────────────────────────────────────────────
    final_report = state.get("final_report") or ""
    if final_report.strip():
        doc.add_page_break()
        doc.add_heading("Full Investment Memo", level=1)
        _add_markdown_report(doc, final_report)

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Confidential — For Internal Use Only")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(output_path)
    return os.path.abspath(output_path)
